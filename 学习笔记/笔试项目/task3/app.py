from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from docx.oxml.ns import qn
from pprint import pprint
app = Flask(__name__)

@app.route('/')
def index():
    # 默认数据，用于前端表单的预填充
    data = {
        "company": "测试企业",
        "legalName": "测试企业法人",
        "legalPhone": "15500000000",
        "employeeInfor": [
            {"num": "1", "name": "张三", "idcard": "4222222222221", "birthday": "19990101"},
            {"num": "2", "name": "李四", "idcard": "4222222222222", "birthday": "19990103"},
            {"num": "3", "name": "王五", "idcard": "4222222222223", "birthday": "19990102"},
            {"num": "4", "name": "赵六", "idcard": "4222222222224", "birthday": "19990104"},
            {"num": "", "name": "刘七", "idcard": "4222222222225", "birthday": "19990105"}
        ]
    }
    return render_template('index.html', data=data)

@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    # 接收并解析前端发送的JSON数据
    data = request.json

    pprint(data) # json数据
    company = data['company']
    legal_name = data['legalName']
    legal_phone = data['legalPhone']
    employees = data['employeeInfor']

    # 创建Word文档
    doc = Document()
    heading = doc.add_heading(level=1)
    heading_run = heading.add_run(f"{company}公司基本信息")
    # 设置标题字体和样式
    heading_run.font.size = Pt(14)
    heading_run.font.name = '仿宋_GB2312'
    heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    heading_run.bold = False

    # 创建表格并填入公司信息
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    table.cell(0, 0).text = '公司名称'
    table.cell(0, 1).text = company
    table.cell(0, 2).text = '公司法人'
    table.cell(0, 3).text = legal_name
    table.cell(1, 0).text = '法人联系方式'
    table.cell(1, 1).text = legal_phone
    table.cell(1, 2).merge(table.cell(1, 3))

    # 设置表格字体
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = '仿宋_GB2312'

    # 添加员工信息标题行
    emp_title_row = table.add_row()
    emp_title_cell = emp_title_row.cells
    emp_title_cell[0].text = '员工信息'
    emp_title_cell[0].merge(emp_title_cell[3])
    for run in emp_title_cell[0].paragraphs[0].runs:
        run.font.size = Pt(12)
        run.font.name = '仿宋_GB2312'

    # 添加每个员工的信息到表格
    for emp in employees:
        emp_num_row = table.add_row()
        emp_num_cell = emp_num_row.cells
        emp_num_cell[0].text = f'员工 {emp.get("num", "")}' if emp.get("num") else '员工'
        emp_num_cell[0].merge(emp_num_cell[3])
        for run in emp_num_cell[0].paragraphs[0].runs:
            run.font.size = Pt(12)
            run.font.name = '仿宋_GB2312'

        emp_name_row = table.add_row()
        emp_name_cell = emp_name_row.cells
        emp_name_cell[0].text = '员工名'
        emp_name_cell[1].text = emp["name"]
        emp_name_cell[1].merge(emp_name_cell[3])
        for cell in emp_name_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = '仿宋_GB2312'

        emp_id_row = table.add_row()
        emp_id_cell = emp_id_row.cells
        emp_id_cell[0].text = '员工身份证'
        emp_id_cell[1].text = emp["idcard"]
        emp_id_cell[1].merge(emp_id_cell[3])
        for cell in emp_id_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = '仿宋_GB2312'

        emp_birthday_row = table.add_row()
        emp_birthday_cell = emp_birthday_row.cells
        emp_birthday_cell[0].text = '出生年月'
        emp_birthday_cell[1].text = emp["birthday"]
        emp_birthday_cell[1].merge(emp_birthday_cell[3])
        for cell in emp_birthday_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = '仿宋_GB2312'

    # 将文档保存到内存中，并准备发送
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    # 发送生成的Word文档
    return send_file(byte_io, as_attachment=True, attachment_filename='公司信息.docx')

if __name__ == '__main__':
    app.run(debug=True)
